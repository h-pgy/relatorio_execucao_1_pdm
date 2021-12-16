from docx import Document
import pandas as pd
from collections import OrderedDict
from .config import GENERATED_DATA_DIR
from .read_data import read_execucao, read_estrutura_pdm
from .utils import solve_path


class Word:

    def __init__(self, file_name = 'word_relatorio_execucao.docx',
                output_folder = GENERATED_DATA_DIR):

        self.read_execucao = read_execucao
        self.read_estrutura = read_estrutura_pdm
        self.doc = Document()
        self.file_path = solve_path(output_folder, file_name)

    def merge_data_estrutura_execucao(self):

        df_execucao = self.read_execucao()
        df_estrutura = self.read_estrutura()

        merged = pd.merge(df_execucao, 
                        df_estrutura, 
                        how = 'left', 
                        left_on = 'Meta', 
                        right_on ='Meta_numero')

        return merged

    def dados_ficha(self, row):
    
        num_meta = row['Meta']
        obs = row['Obs.']
        if not pd.isnull(obs):
            obs = f' {obs}'
        else:
            obs = ''
        
        parsed = OrderedDict({
        'Eixo' : row['Eixo'],
        'Objetivo estratégico' : row['Objetivo Estratégico'],
        f'Meta {num_meta}{obs}' : row['Meta_descricao'],
        'Indicador' : row['Indicador'],
        'Secretarias responsáveis' : row['Secretaria Responsável'],
        'ODS vinculados' : row['ODS vinculados'],
        'Resultado apurado no período' : row['Realizado acumulado']
        })
        
        
        return parsed

    def write_page(self, dados):

        for chave, valor in dados.items():
    
            chave = chave + ':'
            valor = ' ' + str(valor)
            p = self.doc.add_paragraph()
            p.add_run(chave).bold = True
            p.add_run(valor)
            
            
        self.doc.add_page_break()

    def write_document(self, file = None, save=True):

        if file is None:
            file = self.file_path

        df = self.merge_data_estrutura_execucao()

        for i, row in df.iterrows():

            dados = self.dados_ficha(row)
            self.write_page(dados)

        if save:
            self.doc.save(file)

    def __call__(self):

        self.write_document()



    